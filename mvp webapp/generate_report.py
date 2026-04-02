"""
SYNTHETIC REPORT GENERATOR
--------------------------
This script generates a container inspection report from structured JSON data.

Requirements:
    pip install python-docx

Usage:
    python generate_report.py
"""

import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_json_data(filename):
    """
    Load the synthetic data from JSON file

    Args:
        filename (str): Path to JSON file

    Returns:
        dict: Parsed JSON data
    """
    print(f"Loading data from {filename}...")
    with open(filename, 'r', encoding='utf-8') as f:
        data = json.load(f)
    print("Data loaded successfully!")
    return data


def add_header_section(doc, data):
    """
    Add the header section with references and date

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    # Add header line with references
    header = doc.add_paragraph()
    header.add_run(f"OUR REF: {data['header']['isa_reference']}").bold = True
    header.add_run("                YOUR REF: ")
    header.add_run(data['header']['principal_reference']).bold = True
    header.add_run("                DATE: ")
    header.add_run(data['header']['report_date']).bold = True

    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()


def add_background_section(doc, data):
    """
    Add the BACKGROUND section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    # Main heading
    doc.add_heading('BACKGROUND', level=1)

    # Subheading: Circumstances Leading to Claim
    doc.add_heading('Circumstances Leading to Claim', level=2)

    # Get grammar switches for readability
    g = data['grammar_switches']

    # Paragraph 1: Shipment basics
    p1 = doc.add_paragraph()
    p1.add_run(
        f"From documentation and information made available, we understand that the subject consignment, "
        f"comprising {data['shipment_details']['number_of_packages']} of {data['shipment_details']['goods_description']}, "
        f"was sold by the Shipper, {data['shipment_details']['shipper_name']}, {data['shipment_details']['shipper_country']} "
        f"to the Consignee, {data['shipment_details']['consignee_name']}, {data['shipment_details']['consignee_country']} "
        f"on {data['shipment_details']['incoterms']}."
    )

    # Paragraph 2: Container gate out
    p2 = doc.add_paragraph()
    p2.add_run(
        f"According to information secured from the Carrier's online tracking, "
        f"{data['container_details']['number_of_containers']} x empty {data['container_details']['container_types']} "
        f"{g['container_singular_plural']}, No. {data['container_details']['container_numbers']}, "
        f"gated out from the terminal at the port of {data['port_location_details']['origin_port_name']}, "
        f"{data['port_location_details']['origin_port_country']} on {data['container_details']['container_gate_out_date']}."
    )

    # Paragraph 3: Container return and carrier receipt
    p3 = doc.add_paragraph()
    p3.add_run(
        f"The {g['container_singular_plural']} {g['was_were']} returned, fully laden, to the port of "
        f"{data['port_location_details']['origin_port_name']} on {data['container_details']['container_return_date']}, "
        f"where {g['it_they']} {g['was_were']} received by the Carrier, {data['carrier_shipping_details']['carrier_name']} "
        f"for further shipment to {data['port_location_details']['discharge_port_name']}, "
        f"{data['port_location_details']['discharge_port_country']} on {data['carrier_shipping_details']['shipment_terms']} terms "
        f"under cover of Bill of Lading No. {data['carrier_shipping_details']['bill_of_lading_number']} "
        f"issued at {data['carrier_shipping_details']['bill_of_lading_issue_place']} "
        f"on {data['carrier_shipping_details']['bill_of_lading_issue_date']}."
    )

    # Paragraph 4: Vessel loading
    p4 = doc.add_paragraph()
    p4.add_run(
        f"The {g['container_singular_plural']} {g['was_were']} shipped on board the carrying vessel, "
        f"M/V \"{data['carrier_shipping_details']['vessel_name']}\" VOY. {data['carrier_shipping_details']['voyage_number']} "
        f"at {data['port_location_details']['origin_port_name']} on {data['carrier_shipping_details']['vessel_loading_date']}."
    )

    # Paragraph 5: Transhipment (conditional)
    if data['transhipment_details']['has_transhipment']:
        p5 = doc.add_paragraph()
        p5.add_run(
            f"The vessel arrived at the transhipment port of {data['transhipment_details']['transhipment_port_name']}, "
            f"{data['transhipment_details']['transhipment_port_country']} on {data['transhipment_details']['transhipment_arrival_date']} "
            f"where the {g['container_singular_plural']} {g['was_were']} discharged later on the same day and then further loaded "
            f"on board the on-carrying vessel, M/V \"{data['transhipment_details']['oncarrying_vessel_name']}\" "
            f"VOY. {data['transhipment_details']['oncarrying_voyage_number']} at {data['transhipment_details']['transhipment_port_name']} "
            f"on {data['transhipment_details']['transhipment_reload_date']}."
        )

    # Paragraph 6: Final discharge
    p6 = doc.add_paragraph()
    p6.add_run(
        f"The vessel arrived at the final discharge port of {data['final_discharge_delivery']['final_discharge_port_name']}, "
        f"{data['final_discharge_delivery']['final_discharge_port_country']} on {data['final_discharge_delivery']['final_port_arrival_date']} "
        f"where the {g['container_singular_plural']} {g['was_were']} subsequently discharged and moved into the CY for "
        f"temporary storage pending collection."
    )

    # Paragraph 7: Delivery
    p7 = doc.add_paragraph()
    p7.add_run(
        f"Following completion of import formalities, the {g['container_singular_plural']} were collected from the terminal "
        f"at the port on {data['final_discharge_delivery']['container_collection_date']} for delivery to the "
        f"{data['final_discharge_delivery']['consignee_delivery_type']} premises, located at "
        f"{data['final_discharge_delivery']['delivery_premises_location']}, {data['final_discharge_delivery']['delivery_city']}, "
        f"arriving on {data['final_discharge_delivery']['delivery_arrival_date']}."
    )

    # Paragraph 8: Damage discovery
    p8 = doc.add_paragraph()
    p8.add_run(
        f"It was reported that at the time of the delivery of the {g['container_singular_plural']}, {g['it_they']} {g['was_were']} "
        f"found to be in an apparent sound condition, with original shipping seal{g['seal_singular_plural']} still intact, however, "
        f"upon opening of the doors of container No. {data['damage_discovery']['damaged_container_number']}, "
        f"the receiving personnel found that {data['damage_discovery']['damage_discovery_narrative']}."
    )

    # Paragraph 9: Following discovery
    p9 = doc.add_paragraph()
    p9.add_run(
        "Following discovery, the Consignee report the matter to concerned parties, as a result of which, "
        "we were requested to attend survey in order to establish nature, extent and cause of any resulting loss."
    )

    # Subheading: Arrangements for Survey
    doc.add_heading('Arrangements for Survey', level=2)

    # Paragraph 10: Survey arrangements
    p10 = doc.add_paragraph()
    p10.add_run(
        f"Following receipt of instructions, we immediately contacted {data['survey_arrangements']['consignee_contact_person']} "
        f"of the Consignee, in order to make necessary arrangements for survey. From discussions, we understood that "
        f"{data['survey_arrangements']['survey_arrangements_discussion']}"
    )

    # Paragraph 11: Survey date
    p11 = doc.add_paragraph()
    p11.add_run(f"Therefore, arrangements were made to attend inspection on {data['survey_arrangements']['survey_attendance_date']}.")


def add_survey_section(doc, data):
    """
    Add the SURVEY section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    # Main heading
    doc.add_heading('SURVEY', level=1)

    # Get grammar switches
    g = data['grammar_switches']

    # Subheading: Description of Goods and Packaging
    doc.add_heading('Description of Goods and Packaging', level=2)

    p1 = doc.add_paragraph()
    p1.add_run(
        f"The goods forming the subject of this claim comprised {data['shipment_details']['number_of_packages']} of "
        f"{data['shipment_details']['goods_description']} stowed in {data['container_details']['number_of_containers']} x empty "
        f"{data['container_details']['container_types']} {g['container_singular_plural']}, No. {data['container_details']['container_numbers']}. "
        f"GW: {data['goods_packaging']['gross_weight_kgs']} KGS NW: {data['goods_packaging']['net_weight_kgs']} KGS."
    )

    p2 = doc.add_paragraph()
    p2.add_run(data['goods_packaging']['packaging_method_description'])

    # Subheading: Condition of Container
    doc.add_heading('Condition of Container', level=2)

    # Conditional: container available or not
    if data['container_condition']['container_available']:
        p3 = doc.add_paragraph()
        p3.add_run(data['container_condition']['container_condition_description'])
    else:
        p3 = doc.add_paragraph()
        p3.add_run(
            f"At the time of our attendance, the {g['container_singular_plural']} had already been devanned and returned to the Carrier. "
            f"However, from discussions with the Consignee, we understand that {data['container_condition']['container_condition_from_consignee']}"
        )

    # Subheading: Condition of Goods
    doc.add_heading('Condition of Goods', level=2)

    p4 = doc.add_paragraph()
    p4.add_run("At the time of attendance, the goods had already been sorted and set aside by the Consignee, pending survey.")

    p5 = doc.add_paragraph()
    p5.add_run(data['goods_condition']['goods_condition_description'])

    # Optional: Testing section
    if data['testing']['testing_performed']:
        doc.add_heading('Temperature / Chemical Testing / Moisture Testing', level=2)

        if data['testing']['temperature_testing_results']:
            p6 = doc.add_paragraph()
            p6.add_run(data['testing']['temperature_testing_results'])

        if data['testing']['chemical_testing_results']:
            p7 = doc.add_paragraph()
            p7.add_run(data['testing']['chemical_testing_results'])

        if data['testing']['moisture_testing_results']:
            p8 = doc.add_paragraph()
            p8.add_run(data['testing']['moisture_testing_results'])


def add_discussions_section(doc, data):
    """
    Add the DISCUSSIONS section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    doc.add_heading('DISCUSSIONS', level=1)

    p1 = doc.add_paragraph()
    p1.add_run(
        f"Following survey, we discussed the Consignee's further intentions in regard to the cargo and were advised that "
        f"{data['discussions']['post_survey_discussions']}"
    )


def add_developments_section(doc, data):
    """
    Add the DEVELOPMENTS section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    doc.add_heading('4.   DEVELOPMENTS', level=1)

    p1 = doc.add_paragraph()
    p1.add_run(
        f"We continued to maintain contact with the Consignee and on {data['developments']['development_date']} "
        f"were advised that {data['developments']['developments_narrative']}"
    )


def add_loss_quantification_section(doc, data):
    """
    Add the QUANTIFICATION OF LOSS section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    doc.add_heading('5.   QUANTIFICATION OF LOSS', level=1)

    # 5.1 Loss
    doc.add_heading('5.1    Loss', level=2)

    p1 = doc.add_paragraph()
    p1.add_run(
        f"According to Commercial Invoice No. {data['loss_quantification']['commercial_invoice_number']} "
        f"dated {data['loss_quantification']['commercial_invoice_date']}, the value of the goods forming the subject "
        f"of this claim amounts to {data['loss_quantification']['claim_value_currency']} "
        f"{data['loss_quantification']['claim_value_amount']} {data['shipment_details']['incoterms']}."
    )

    p2 = doc.add_paragraph()
    p2.add_run(data['loss_quantification']['loss_details_narrative'])

    # 5.2 Additional Costs
    doc.add_heading('5.2    Additional Costs', level=2)

    p3 = doc.add_paragraph()
    p3.add_run(data['loss_quantification']['additional_costs_narrative'])


def add_cause_of_loss_section(doc, data):
    """
    Add the CAUSE OF LOSS section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    doc.add_heading('6.   CAUSE OF LOSS', level=1)

    p1 = doc.add_paragraph()
    p1.add_run(
        f"From findings during survey, we attribute the loss in this instance to "
        f"{data['cause_of_loss']['loss_cause_summary']}."
    )

    p2 = doc.add_paragraph()
    p2.add_run(data['cause_of_loss']['loss_cause_explanation'])


def add_photographs_section(doc, data):
    """
    Add the PHOTOGRAPHS section

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    doc.add_heading('7.   PHOTOGRAPHS', level=1)

    p1 = doc.add_paragraph()
    p1.add_run(
        "Photographs taken at the time of survey, along with those supplied by parties concerned are embedded "
        "within the body of this report."
    )

    # Add closing statement
    p2 = doc.add_paragraph()
    p2.add_run(
        "This Certificate of Survey is issued, without prejudice, and subject to the terms and conditions "
        "of the relative Policy of Insurance."
    )


def add_footer_section(doc, data):
    """
    Add the footer with signature block

    Args:
        doc: Word document object
        data: JSON data dictionary
    """
    doc.add_paragraph()
    doc.add_paragraph()

    # "for" line
    p1 = doc.add_paragraph()
    p1.add_run("for")

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature line
    p2 = doc.add_paragraph()
    p2.add_run("SURVEYOR")
    doc.add_paragraph("__________")

    doc.add_paragraph()
    doc.add_paragraph()

    # Documents enclosed
    p3 = doc.add_paragraph()
    p3.add_run("Copy documents enclosed:")


def generate_report(json_filename, output_filename):
    """
    Main function to generate the complete report

    Args:
        json_filename (str): Path to input JSON file
        output_filename (str): Path to output Word document
    """
    print("=" * 60)
    print("CONTAINER INSPECTION REPORT GENERATOR")
    print("=" * 60)
    print()

    # Step 1: Load the JSON data
    data = load_json_data(json_filename)

    # Step 2: Create a new Word document
    print("Creating Word document...")
    doc = Document()

    # Step 3: Add all sections in order
    print("Adding HEADER section...")
    add_header_section(doc, data)

    print("Adding BACKGROUND section...")
    add_background_section(doc, data)

    print("Adding SURVEY section...")
    add_survey_section(doc, data)

    print("Adding DISCUSSIONS section...")
    add_discussions_section(doc, data)

    print("Adding DEVELOPMENTS section...")
    add_developments_section(doc, data)

    print("Adding QUANTIFICATION OF LOSS section...")
    add_loss_quantification_section(doc, data)

    print("Adding CAUSE OF LOSS section...")
    add_cause_of_loss_section(doc, data)

    print("Adding PHOTOGRAPHS section...")
    add_photographs_section(doc, data)

    print("Adding FOOTER section...")
    add_footer_section(doc, data)

    # Step 4: Save the document
    print(f"Saving report to {output_filename}...")
    doc.save(output_filename)

    print()
    print("=" * 60)
    print("REPORT GENERATED SUCCESSFULLY!")
    print("=" * 60)
    print(f"Output file: {output_filename}")
    print(f"Case reference: {data['header']['isa_reference']}")
    print()


if __name__ == "__main__":
    # Configuration
    INPUT_JSON = "synthetic_data.json"
    OUTPUT_DOCX = "Generated_Synthetic_Report.docx"

    # Generate the report
    generate_report(INPUT_JSON, OUTPUT_DOCX)
