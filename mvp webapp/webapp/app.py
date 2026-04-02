"""
CONTAINER INSPECTION REPORT WEB APPLICATION
Flask Backend Server
"""

from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from document_extractor import process_uploaded_documents

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'  # Change this in production!

# Configuration
UPLOAD_FOLDER = 'uploads'
REPORTS_FOLDER = 'reports'
DATA_FOLDER = 'data'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'jpg', 'jpeg', 'png'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORTS_FOLDER'] = REPORTS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create necessary directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)
os.makedirs(DATA_FOLDER, exist_ok=True)

# Simple user database (in production, use proper database)
USERS = {
    'admin': 'admin123',
    'surveyor': 'survey123',
    'demo': 'demo123'
}


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def load_case_data():
    """Load all case data from JSON files"""
    cases_file = os.path.join(DATA_FOLDER, 'cases.json')
    if os.path.exists(cases_file):
        with open(cases_file, 'r') as f:
            return json.load(f)
    return []


def save_case_data(cases):
    """Save case data to JSON file"""
    cases_file = os.path.join(DATA_FOLDER, 'cases.json')
    with open(cases_file, 'w') as f:
        json.dump(cases, f, indent=2)


def generate_simple_report(case_data):
    """Generate a simple Word report from case data"""
    doc = Document()

    # Header with Company Branding
    header = doc.add_paragraph()
    header.add_run(f"QUEST MARINE LLC\n").bold = True
    header.add_run(f"Marine Surveyors and Consultants\n")
    header.add_run(f"ISO 9001, 14001, 45001 Certified Company\n\n")

    header.add_run(f"CONTAINER INSPECTION REPORT\n").bold = True
    header.add_run(f"Case Reference: {case_data.get('case_reference', 'N/A')}\n")
    header.add_run(f"Date: {case_data.get('report_date', datetime.now().strftime('%d %B %Y'))}\n")

    doc.add_paragraph()

    # Basic Information
    doc.add_heading('CASE INFORMATION', level=1)
    p1 = doc.add_paragraph()
    p1.add_run(f"Container Number: {case_data.get('container_number', 'N/A')}\n")
    p1.add_run(f"B/L Number: {case_data.get('bl_number', 'N/A')}\n")
    p1.add_run(f"Goods Description: {case_data.get('goods_description', 'N/A')}\n")
    p1.add_run(f"Shipper: {case_data.get('shipper', 'N/A')}\n")
    p1.add_run(f"Consignee: {case_data.get('consignee', 'N/A')}\n")

    doc.add_paragraph()

    # Documents Uploaded
    doc.add_heading('DOCUMENTS RECEIVED', level=1)
    docs = case_data.get('documents', {})
    p2 = doc.add_paragraph()
    p2.add_run(f"Bill of Lading: {'✓' if docs.get('bill_of_lading') else '✗'}\n")
    p2.add_run(f"Commercial Invoice: {'✓' if docs.get('commercial_invoice') else '✗'}\n")
    p2.add_run(f"Packing List: {'✓' if docs.get('packing_list') else '✗'}\n")
    p2.add_run(f"iAuditor/Safety Culture Report: {'✓' if docs.get('iauditor_report') else '✗'}\n")

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.add_run("\n\n_" * 40 + "\n")
    footer.add_run("QUEST MARINE LLC\n").bold = True
    footer.add_run("Marine Surveyors and Consultants\n")
    footer.add_run("ISO 9001, 14001, 45001 Certified Company\n\n")
    footer.add_run(f"Report Generated on: {datetime.now().strftime('%d %B %Y %H:%M:%S')}")

    # Save the document
    filename = f"Report_{case_data.get('case_reference', 'CASE')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_FOLDER, filename)
    doc.save(filepath)

    return filename


# ===========================
# ROUTES
# ===========================

@app.route('/')
def index():
    """Redirect to login page"""
    if 'username' in session:
        return redirect(url_for('upload'))
    return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page"""
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        if username in USERS and USERS[username] == password:
            session['username'] = username
            flash('Login successful!', 'success')
            return redirect(url_for('upload'))
        else:
            flash('Invalid username or password', 'error')

    return render_template('login.html')


@app.route('/logout')
def logout():
    """Logout user"""
    session.pop('username', None)
    flash('You have been logged out', 'info')
    return redirect(url_for('login'))


@app.route('/upload', methods=['GET', 'POST'])
def upload():
    """Document upload page"""
    if 'username' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        # Get case information
        case_reference = request.form.get('case_reference', f'CASE-{datetime.now().strftime("%Y%m%d-%H%M%S")}')
        container_number = request.form.get('container_number', '')
        bl_number = request.form.get('bl_number', '')
        goods_description = request.form.get('goods_description', '')
        shipper = request.form.get('shipper', '')
        consignee = request.form.get('consignee', '')

        # Create case folder
        case_folder = os.path.join(UPLOAD_FOLDER, case_reference)
        os.makedirs(case_folder, exist_ok=True)

        # Handle file uploads
        documents = {}
        file_types = ['bill_of_lading', 'commercial_invoice', 'packing_list', 'iauditor_report']

        for file_type in file_types:
            if file_type in request.files:
                file = request.files[file_type]
                if file and file.filename and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(case_folder, f"{file_type}_{filename}")
                    file.save(filepath)
                    documents[file_type] = filename

        # Save case data
        case_data = {
            'case_reference': case_reference,
            'container_number': container_number,
            'bl_number': bl_number,
            'goods_description': goods_description,
            'shipper': shipper,
            'consignee': consignee,
            'documents': documents,
            'created_by': session['username'],
            'created_at': datetime.now().isoformat(),
            'status': 'pending'
        }

        # Load existing cases and add new one
        cases = load_case_data()
        cases.append(case_data)
        save_case_data(cases)

        flash(f'Case {case_reference} created successfully!', 'success')
        return redirect(url_for('generate_report_page', case_ref=case_reference))

    return render_template('upload.html', username=session['username'])


@app.route('/generate/<case_ref>')
def generate_report_page(case_ref):
    """Page to generate report"""
    if 'username' not in session:
        return redirect(url_for('login'))

    # Find case data
    cases = load_case_data()
    case_data = next((c for c in cases if c['case_reference'] == case_ref), None)

    if not case_data:
        flash('Case not found', 'error')
        return redirect(url_for('upload'))

    return render_template('generate.html', case=case_data, username=session['username'])


@app.route('/api/generate-report/<case_ref>', methods=['POST'])
def api_generate_report(case_ref):
    """API endpoint to generate report"""
    if 'username' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    # Find case data
    cases = load_case_data()
    case_data = next((c for c in cases if c['case_reference'] == case_ref), None)

    if not case_data:
        return jsonify({'error': 'Case not found'}), 404

    # Generate report
    try:
        case_data['report_date'] = datetime.now().strftime('%d %B %Y')
        filename = generate_simple_report(case_data)

        # Update case status
        case_data['status'] = 'completed'
        case_data['report_file'] = filename
        case_data['completed_at'] = datetime.now().isoformat()
        save_case_data(cases)

        return jsonify({
            'success': True,
            'message': 'Report generated successfully',
            'filename': filename,
            'download_url': url_for('download_report', filename=filename)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download/<filename>')
def download_report(filename):
    """Download generated report"""
    if 'username' not in session:
        return redirect(url_for('login'))

    filepath = os.path.join(REPORTS_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        flash('Report not found', 'error')
        return redirect(url_for('dashboard'))


@app.route('/dashboard')
def dashboard():
    """Dashboard page showing all cases"""
    if 'username' not in session:
        return redirect(url_for('login'))

    cases = load_case_data()

    # Sort by created date (newest first)
    cases.sort(key=lambda x: x.get('created_at', ''), reverse=True)

    # Calculate statistics
    total_cases = len(cases)
    pending_cases = len([c for c in cases if c.get('status') == 'pending'])
    completed_cases = len([c for c in cases if c.get('status') == 'completed'])

    stats = {
        'total': total_cases,
        'pending': pending_cases,
        'completed': completed_cases
    }

    return render_template('dashboard.html', cases=cases, stats=stats, username=session['username'])


@app.route('/api/extract-data', methods=['POST'])
def api_extract_data():
    """API endpoint to extract data from uploaded documents"""
    if 'username' not in session:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        # Get uploaded files
        documents = {}
        file_types = ['bill_of_lading', 'commercial_invoice', 'packing_list', 'iauditor_report']
        temp_folder = os.path.join(UPLOAD_FOLDER, 'temp_extraction')
        os.makedirs(temp_folder, exist_ok=True)

        # Save files temporarily
        for file_type in file_types:
            if file_type in request.files:
                file = request.files[file_type]
                if file and file.filename and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(temp_folder, f"{file_type}_{filename}")
                    file.save(filepath)
                    documents[file_type] = filename

        if not documents:
            return jsonify({'error': 'No valid documents uploaded', 'extracted': False}), 400

        # Process documents and extract data
        extracted_data = process_uploaded_documents(temp_folder, documents)

        # Clean up temp files
        import shutil
        if os.path.exists(temp_folder):
            shutil.rmtree(temp_folder)

        return jsonify(extracted_data)

    except Exception as e:
        return jsonify({'error': f'Extraction failed: {str(e)}', 'extracted': False}), 500


@app.route('/contact-admin')
def contact_admin():
    """Contact admin page"""
    return render_template('contact.html')


if __name__ == '__main__':
    print("\n" + "=" * 60)
    print("CONTAINER INSPECTION REPORT SYSTEM")
    print("=" * 60)
    print("\nDefault Login Credentials:")
    print("  Username: admin    Password: admin123")
    print("  Username: surveyor Password: survey123")
    print("  Username: demo     Password: demo123")
    print("\nServer starting...")
    print("Access the application at: http://localhost:5001")
    print("=" * 60 + "\n")

    app.run(debug=True, host='0.0.0.0', port=5001)
